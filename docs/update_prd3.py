import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = docx.Document()

# Styles
styles = doc.styles

# Heading 1 (Green)
h1_style = styles['Heading 1']
h1_font = h1_style.font
h1_font.color.rgb = RGBColor(0x1B, 0x5E, 0x20) # Dark Green
h1_font.name = 'Arial'

# Heading 2 (Green)
h2_style = styles['Heading 2']
h2_font = h2_style.font
h2_font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Green
h2_font.name = 'Arial'

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(text, bold=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    return p

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('CAR DA GENTE\n')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20) # Green
run.font.name = 'Arial'

run2 = title_p.add_run('O CARinho Tecnológico de SimplifiCAR o Cadastro Ambiental Rural\n')
run2.bold = True
run2.italic = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
run2.font.name = 'Arial'

run3 = title_p.add_run('PRD — Desafio 3 | haCARthon 2026 | Versão 4.0 | Junho/2026')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
run3.font.name = 'Arial'

doc.add_paragraph()

# Table Meta
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
# Set background color for cells
for row in table.rows:
    for cell in row.cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="D4EDDA"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

data = [
    ('Projeto', 'CAR da Gente'),
    ('Desafio', 'Desafio 3 — haCARthon 2026'),
    ('Foco', 'Aumentar o Entendimento das Legislações do CAR'),
    ('Versão do PRD', '4.0 — Junho/2026'),
    ('Personas Atendidas', 'Seu Raimundo (produtor rural) · Liderança Comunitária'),
    ('Arquitetura Core', 'IA Assistida + Motor Determinístico + n8n + WhatsApp + MarCARtplace')
]

for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x00)
    table.rows[i].cells[1].text = v
    table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x00)

doc.add_page_break()

# Sumário (TOC)
add_heading('Sumário', level=1)
add_paragraph('1. Visão Geral do Produto')
add_paragraph('2. Personas Centrais')
add_paragraph('3. Objetivos Estratégicos')
add_paragraph('4. Arquitetura de Canais & Integração n8n')
add_paragraph('5. Módulos e Funcionalidades Técnicas')
add_paragraph('   5.1 Central do Produtor')
add_paragraph('   5.2 ComuniCAR: Engajamento Comunitário')
add_paragraph('6. Arquitetura de Segurança e Confiança')
add_paragraph('7. Comparativo Técnico — Antes vs. Depois')
add_paragraph('8. Design Centrado no Usuário (DCU) — Nomenclaturas')
add_paragraph('9. Estudo de Viabilidade Financeira (Calibrado para MVP/Hackathon)')
add_paragraph('10. Considerações Finais')

doc.add_page_break()

# 1. Visão Geral
add_heading('1. Visão Geral do Produto')
add_paragraph('O CAR da Gente é uma plataforma focada na simplificação da regularização ambiental rural no Brasil. Por meio de uma arquitetura de Inteligência Artificial assistida orquestrada pela "CARlinha", a plataforma:')
add_paragraph('Automatiza a triagem técnica e a análise inicial de documentos e fotos de campo', style='List Bullet')
add_paragraph('Traduz o "juridiquês" ambiental em orientações claras para o pequeno produtor rural', style='List Bullet')
add_paragraph('Opera perfeitamente de forma mobile-first, sendo o meio preferido de acesso do produtor via WhatsApp e WebApp leve', style='List Bullet')
add_paragraph('Produz conteúdo educativo viral focado nas comunidades rurais via módulo CarCantado', style='List Bullet')
add_paragraph('Disponibiliza o MarCARtplace: portal de comercialização de Cotas de Reserva Ambiental (CRA) gerando renda e conexão direta', style='List Bullet')

# 2. Personas Centrais
add_heading('2. Personas Centrais')
add_paragraph('Produtor Rural Familiar (Persona Primária)', bold=True)
add_paragraph('Perfil: Baixo letramento digital e escolaridade básica. Possui pequena propriedade que depende de crédito subsidiado (Pronaf) e é refém de despachantes caros. Usa o celular exclusivamente para WhatsApp, redes sociais e vídeos curtos.')
add_paragraph('Dor: Medo de multas do órgão ambiental. Dificuldade extrema de compreender o jargão do Código Florestal e os relatórios do CAR. Necessita orientações diretas e provas visuais do que precisa ser corrigido.')

# 3. Objetivos Estratégicos
add_heading('3. Objetivos Estratégicos')
add_paragraph('Simplificação e Acessibilidade: Eliminar barreiras linguísticas e tecnológicas, com interface voltada 100% à experiência e visão do produtor no canal de maior alcance do país (WhatsApp).', style='List Bullet')
add_paragraph('Arquitetura Segura Assistida: IA restrita a orientações sem substituir o fluxo oficial ou criar falsas regras ambientais. Cálculos baseados no motor determinístico embutido.', style='List Bullet')
add_paragraph('Viralidade e Pertencimento: Transformar a burocracia ambiental em conteúdo engajador através do CarCantado, estimulando a adoção orgânica (UGC).', style='List Bullet')
add_paragraph('Mercado Verde na Palma da Mão: O MarCARtplace como motor financeiro para pequenos produtores rentabilizarem excessos de vegetação nativa.', style='List Bullet')

# 4. Arquitetura de Canais
add_heading('4. Arquitetura de Canais & Integração n8n')
add_heading('4.1 Engine n8n (Canal WhatsApp)', level=2)
add_paragraph('Orquestra os webhooks do WhatsApp, recebe imagens, PDFs e áudios do produtor, aciona os agentes de IA correspondentes e responde em segundos com dados mastigados e linguagem acolhedora. É o coração da acessibilidade rural do sistema.')
add_heading('4.2 Plataforma Web Corporativa', level=2)
add_paragraph('Interface rica para visualização de mapas geoespaciais, controle de metas comunitárias (Termômetro Verde) e geração de minutas automatizadas para o produtor.')
add_heading('4.3 WebApp Responsivo (Campo)', level=2)
add_paragraph('Versão leve e otimizada para redes rurais instáveis. Permite que o produtor navegue pelo mapa do CAR com GPS ativo, visualize APPs e Reservas Legais em sobreposição real e consulte pendências sem precisar de conexão estável.')
add_heading('4.4 Triagem e Análise de Imagem por IA', level=2)
add_paragraph('Ao enviar fotos de evidência de campo, um motor visual ajuda a classificar e confirmar o que há no terreno (APP, Reserva, Área de Uso), orientando o produtor se a foto serve para sanar pendências, apoiado pela CARlinha.')

# 5. Módulos e Funcionalidades Técnicas
add_heading('5. Módulos e Funcionalidades Técnicas')
add_heading('5.1 MÓDULO I — Central do Produtor', level=2)

add_paragraph('RF01 - Juridiquês junto ao Processo — Laudo Cidadão', bold=True)
add_paragraph('Dentro dos documentos pendentes haverá um botão "Quero entender" — Laudo Cidadão. Motor determinístico e auditável que codifica o Código Florestal (Arts. 4, 12, 61-A/B/C, a "escadinha", tetos de RL) como regras executáveis. Ingere a geometria do CAR e cruza dados para gerar um laudo personalizado (déficit, faixa de recomposição, passo a passo).')
add_paragraph('🃏 CARta na Manga 1 — Escadinha Automática para Pequenos Produtores (Art. 61-A): O sistema calcula instantaneamente o tamanho reduzido da recomposição de APP.', style='List Bullet')
add_paragraph('🃏 CARta na Manga 2 — Integração com Mercado de CRA (Art. 66): O sistema sugere "Você tem 5 hectares de déficit. Clique aqui para compensar" direcionando ao MarCARtplace.', style='List Bullet')
add_paragraph('🃏 CARta na Manga 3 — Geração Inteligente do Termo de Compromisso: Gera o Termo pronto para assinar.', style='List Bullet')
add_paragraph('Benefício: Produtor recebe um documento personalizado em linguagem simples que traduz EXATAMENTE a lei, sem intermediários.')

add_paragraph('RF02 - Seu Mundinho — Integração ao WhatsApp (n8n Engine)', bold=True)
add_paragraph('O produtor tira foto da notificação ou envia o número do recibo no WhatsApp. O n8n captura, consulta o motor e devolve explicação acolhedora com citação legal. Exemplo: "Calma, Raimundinho! Veja o que podemos fazer." e conecta às CARtas na manga.')

add_paragraph('RF03 - Calculadora Financeira — Simulador de Incentivos', bold=True)
add_paragraph('Interface que calcula o "Antes vs. Depois" financeiro. Mostra o quanto se perde por linhas de crédito bloqueadas e o quanto ganha com o Selo CARimbo Verde (juros menores, benefícios ambientais). Engajamento puramente econômico.')

add_paragraph('RF04 - Selo CARimbo Verde — A Conquista da Regularização', bold=True)
add_paragraph('Após a regularização, o produtor recebe um selo comemorativo via WhatsApp ("Bolão na regularização!"). Ele emite o selo, compartilha e pode gerar uma animação de 30s da sua história, criando prova social.')

add_heading('5.2 MÓDULO II — ComuniCAR: Engajamento Comunitário', level=2)

add_paragraph('RF05 - CAR Cantado — Educação Ambiental em Ritmo Rural', bold=True)
add_paragraph('Vídeos educativos com o Código Florestal em paródias musicais (piseiro, sertanejo). Distribuídos via TikTok, WhatsApp, e áudios MP3 para Rádios Comunitárias. Fixa o conhecimento via música e cultura regional.')

add_paragraph('RF06 - TiCARtok — Canal de Divulgação', bold=True)
add_paragraph('Canal multiplataforma agregando especialistas, jovens estudantes agro e o CAR Cantado, para combater desinformação com conteúdo responsável e divertido.')

add_paragraph('RF07 - MarCARtplace — Mercado de CRA', bold=True)
add_paragraph('Hub transacional que conecta quem tem excedente de mata a quem precisa compensar. Botões e banners em toda a plataforma incentivam: "Ei, produtor! Você tem excedente de vegetação nativa? Divulgue aqui." Transformando a regularização em lucro real.')

# 6. Arquitetura de Segurança
add_heading('6. Arquitetura de Segurança e Confiança')
add_paragraph('🛡️ Filtro Determinístico Anti-Alucinação', bold=True)
add_paragraph('Camada de software que atua como barreira inviolável entre os LLMs e as bases de dados oficiais. Todo número de área, limite geográfico e enquadramento legal é calculado por fórmulas exatas em código tradicional auditável. A IA generativa é proibida de gerar dados técnicos — sua única função é ler o resultado determinístico e reescrevê-lo em linguagem acolhedora. Resultado: Segurança absoluta. A IA nunca inventa uma lei ou área.')

# 7. Comparativo
add_heading('7. Comparativo Técnico — Antes vs. Depois')
table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Critério'
hdr_cells[1].text = 'SICAR Atual'
hdr_cells[2].text = 'CAR da Gente'

rows_data1 = [
    ('Linguagem Técnico-Jurídica', 'Notificações brutas com termos pesados e siglas.', 'Tradução automatizada via LLM com citações legais — "Laudo Cidadão".'),
    ('Canais de Acesso', 'Exclusivamente via navegador Web Desktop, pesado.', 'Omnichannel: WebApp responsivo leve e canal interativo no WhatsApp.'),
    ('Engajamento e Cultura', 'Obrigatoriedade sob pena de multa ou bloqueio.', 'Incentivos financeiros e comunicação viral (CarCantado / TiCARtok).'),
    ('Acesso Offline / Rural', 'Sem suporte.', 'Spots para rádios comunitárias; WebApp otimizado.'),
    ('Participação Comunitária', 'Individual; sem mecanismo de engajamento.', 'MarCARtplace: mercado verde comunitário.'),
    ('Mercado Ambiental', 'Sem integração direta com CRA.', 'Conecta excedente e déficit com clique direto do Laudo.')
]
for crit, antes, depois in rows_data1:
    row = table1.add_row().cells
    row[0].text = crit
    row[1].text = antes
    row[2].text = depois

# 8. DCU
add_heading('8. Design Centrado no Usuário (DCU) — Nomenclaturas')
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Interface / Ator'
hdr_cells2[1].text = 'Nomenclatura Genérica (antes)'
hdr_cells2[2].text = 'Nomenclatura DCU (CAR da Gente)'

rows_data2 = [
    ('Produtor (WhatsApp)', 'Enviar arquivo / Upload', '📄 Traduzir minha Notificação'),
    ('Produtor (WhatsApp)', 'Desbloquear / Resolver', '🔓 Desatar Pendência'),
    ('Produtor (WebApp – GPS)', 'Ver Mapa / Ativar GPS', '🗺️ Ver no Chão da Terra'),
    ('Produtor (Simulador)', 'Calcular Crédito / Simular', '💰 Destravar meu Crédito Verde'),
    ('Produtor (Regularizado)', 'Emitir Certidão / Baixar PDF', '🏅 Receber meu CARimbo Verde'),
    ('Liderança Comunitária', 'Download MP3', '📻 Baixar para o Radinho da Terra'),
    ('Liderança Comunitária', 'Gravar Depoimento', '🗣️ Contar minha História da Terra'),
    ('Produtor (Excedente)', 'Cadastrar área / Oferecer CRA', '🌿 VENDA AQUI SUA COTA')
]
for ator, antes, depois in rows_data2:
    row = table2.add_row().cells
    row[0].text = ator
    row[1].text = antes
    row[2].text = depois

# 9. Viabilidade
add_heading('9. Estudo de Viabilidade Financeira (Calibrado para MVP)')
add_paragraph('O plano operacional assume uma validação inicial de 12 meses focada no desenvolvimento do produto base, cortando licenças desnecessárias e concentrando esforços em freelancers super enxutos.')

add_heading('9.1 Investimento Total (Ano 1)', level=2)
add_paragraph('Total Estimado: R$ 86.820')
add_paragraph('Tecnologia & Infraestrutura (55% — R$ 47.820): Desenvolvimento Full-Stack (part-time), custos mínimos de API de IA (LLMs e Visão), Supabase.', style='List Bullet')
add_paragraph('Módulo CarCantado (20% — R$ 17.000): Apenas Videomaker Freelancer para edição base (R$ 1.200/mês) + ferramentas (CapCut Pro).', style='List Bullet')
add_paragraph('Produto & Operações (15% — R$ 13.000): UX/UI acessível freelance, marketing pontual e abertura CNPJ.', style='List Bullet')
add_paragraph('Reserva de Contingência (10% — R$ 9.000).', style='List Bullet')

add_heading('9.2 Indicadores Financeiros e Projeção', level=2)
add_paragraph('Apesar dos cortes drásticos, a escalabilidade orgânica foca no Módulo CarCantado e no MarCARtplace.')
add_paragraph('Meta de Adoção (Ano 1): 2.000 usuários ativos e 300 pagantes.', style='List Bullet')
add_paragraph('Receita Y1: R$ 115.000 (SaaS freemium, relatórios).', style='List Bullet')
add_paragraph('Break-even Projetado: 14 meses devido aos custos extremamente enxutos.', style='List Bullet')
add_paragraph('ROI Estimado em 36 meses: 180%.', style='List Bullet')

add_heading('10. Considerações Finais')
add_paragraph('O CAR da Gente agora opera em "modo guerrilha", utilizando o mínimo de recursos operacionais para validar o engajamento rural através de inteligência artificial de baixo custo e vídeos virais, promovendo inclusão sem burocracia e com custos de hackathon muito enxutos.')

# If the file is open, it might fail to save to Final again, so let's use a unique name
doc.save('docs/CAR_da_Gente_Final_v2.docx')
